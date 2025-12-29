import streamlit as st
import fitz  # PyMuPDF
import google.generativeai as genai
from PIL import Image
import io
import re
import subprocess
import sys
import glob
import os
import json
import time
import base64
import uuid

# Robust Import for Word generation
try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None
    Inches = None

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="PhD Agent: Experiment Designer", layout="wide")
st.title("PhD Assistant: Agentic Experiment Designer 🧪")

# --- 2. SIDEBAR ---
with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = st.text_input("1. Enter Google API Key", type="password")
    
    selected_model_name = None
    if api_key:
        try:
            genai.configure(api_key=api_key)
            models = list(genai.list_models())
            
            # Filter specifically for Flash or Flash-Lite models
            valid_models = [
                m.name for m in models 
                if 'generateContent' in m.supported_generation_methods 
                and 'flash' in m.name.lower()
            ]
            
            # Sort to likely show newest versions first
            valid_models.sort(reverse=True)
            
            if valid_models:
                selected_model_name = st.selectbox("Select Model", valid_models, index=0)
            else:
                st.warning("No Flash models found. Using fallback.")
                selected_model_name = "models/gemini-2.5-flash-lite" # Fallback
                
        except Exception as e:
            st.error(f"API Key Error: {e}")

    st.divider()
    
    # --- CASE MANAGEMENT ---
    st.header("💾 Case Management")
    HISTORY_FILE = "experiment_history.json"
    
    # Init History
    if "history" not in st.session_state:
        if os.path.exists(HISTORY_FILE):
            try:
                with open(HISTORY_FILE, "r") as f: st.session_state.history = json.load(f)
            except: st.session_state.history = {}
        else: st.session_state.history = {}
    
    if "current_case_name" not in st.session_state:
        st.session_state.current_case_name = None

    # Helper Functions for Serialization
    def serialize_run_output(run_output):
        if not run_output: return None
        return {
            "out": run_output["out"],
            "err": run_output["err"],
            "plots": [base64.b64encode(p).decode('utf-8') for p in run_output["plots"]]
        }

    def deserialize_run_output(run_data):
        if not run_data: return None
        return {
            "out": run_data.get("out", ""),
            "err": run_data.get("err", ""),
            "plots": [base64.b64decode(p) for p in run_data.get("plots", [])]
        }

    # 1. Create New Case
    save_name = st.text_input("New Case Name", placeholder="e.g. Test 1", label_visibility="collapsed")
    if st.button("💾 Create New Case"):
        if save_name:
            st.session_state.history[save_name] = {
                "extracted_text": st.session_state.get("extracted_text", ""),
                "architect_plan": st.session_state.get("architect_plan", ""),
                "active_code": st.session_state.get("active_code", ""),
                "run_output": serialize_run_output(st.session_state.get("run_output", None))
            }
            st.session_state.current_case_name = save_name
            with open(HISTORY_FILE, "w") as f: json.dump(st.session_state.history, f)
            st.success(f"Created '{save_name}'!")
            st.rerun()

    # 2. Update Existing Case
    if st.session_state.current_case_name:
        st.caption(f"Editing: **{st.session_state.current_case_name}**")
        if st.button("💾 Save All (Sidebar)", type="secondary"):
            st.session_state.history[st.session_state.current_case_name] = {
                "extracted_text": st.session_state.get("extracted_text", ""),
                "architect_plan": st.session_state.get("architect_plan", ""),
                "active_code": st.session_state.get("active_code", ""),
                "run_output": serialize_run_output(st.session_state.get("run_output", None))
            }
            with open(HISTORY_FILE, "w") as f: json.dump(st.session_state.history, f)
            st.toast(f"Saved '{st.session_state.current_case_name}'!")

    # 3. Load Case
    if st.session_state.history:
        load_name = st.selectbox("Load Case", list(st.session_state.history.keys()))
        if st.button("📂 Load Selected"):
            data = st.session_state.history[load_name]
            st.session_state.current_case_name = load_name
            st.session_state.extracted_text = data.get("extracted_text", "")
            st.session_state.architect_plan = data.get("architect_plan", "")
            st.session_state.active_code = data.get("active_code", "")
            st.session_state.run_output = deserialize_run_output(data.get("run_output"))
            st.session_state.show_review = True
            st.rerun()

    st.divider()
    uploaded_file = st.file_uploader("3. Upload Notes (PDF)", type="pdf")

# --- 3. THE AGENTIC ENGINE ---

def extract_code(text):
    match = re.search(r"```(python|py)?\n(.*?)```", text, re.DOTALL)
    if match:
        code = match.group(2).strip()
        # SANITIZER: Hard-fix syntax hallucinations
        code = code.replace("<<=", "<<").replace(">>=", ">>")
        return code
    return None

def run_architect_agent(problem, model_name):
    """Step 1: Generates the Math Plan"""
    model = genai.GenerativeModel(model_name)
    with st.spinner("🧠 Architect is planning the logic..."):
        architect_prompt = f"""
        Analyze this Control Theory problem:
        {problem}
        
        TASK:
        1. Identify the dimensions of all matrices (A, B, F, E_A, etc.).
           - LOOK FOR CLUES: "I" usually means Identity (Square Matrix). "0" implies dimensions matching neighbors.
           - CHECK COMPATIBILITY: If M = p*I and I is 2x2, then M must be 2x2 (not a vector).
        2. Map the LMI block structure.
        3. Identify which variables are scalars (1x1) vs matrices.
        4. Explicitly state the multiplication rule: "Scalar (1x1) * Matrix" uses `*`, "Matrix * Matrix" uses `@`.
        
        Return a concise math plan (Plain text/Markdown).
        """
        # FIX: Low temperature for stable reasoning
        plan_resp = model.generate_content(architect_prompt, generation_config={"temperature": 0.1})
        return plan_resp.text

def run_coder_and_debugger(plan, template, model_name):
    """Step 2 & 3: Generates Code from Plan and Debugs it"""
    model = genai.GenerativeModel(model_name)
    
    # STEP 2: THE CODER (Implementation)
    with st.status("💻 Coder: Translating plan to Python...", expanded=False) as status:
        coder_prompt = f"""
        Using this Math Plan: {plan}
        And this Golden Template Style: {template}
        
        Write the complete Python code. 
        - Use * for scalars, @ for matrices.
        - Use cp.bmat for the LMI.
        - Ensure numerical stability: Add small epsilon (1e-6) to constraints if needed (e.g. nu >= 1e-6).
        - CRITICAL: Use raw strings (r"...") for ALL Matplotlib text containing backslashes or LaTeX. 
          Correct: plt.xlabel(r"Value of $\\rho$")
          Wrong: plt.xlabel("Value of $\\rho$")
        - Solver Robustness: Use a loop [cp.CLARABEL, cp.SCS, cp.ECOS] to handle "inaccurate solution" warnings.
        
        SOLVER SPECIFIC RULES:
        1. CLARABEL does NOT accept 'eps'. Do not pass it.
        2. When printing errors, do NOT use `solver.name` (solver is likely a string). Use `str(solver)` or `solver`.
        
        Return ONLY the code block.
        """
        # FIX: Zero temperature for deterministic code generation
        code_resp = model.generate_content(coder_prompt, generation_config={"temperature": 0.0})
        raw_code = extract_code(code_resp.text)
        status.update(label="💻 Coder: Draft complete!", state="complete")

    # STEP 3: THE DEBUGGER (Verification)
    with st.status("🕵️ Debugger: Pre-flight Syntax Check...", expanded=True) as status:
        debugger_prompt = f"""
        Review this generated code for errors:
        {raw_code}
        
        Check for:
        1. Invalid syntax like `<<=` (must be `<< 0`).
        2. Dimension mismatches in `cp.bmat`.
        3. Incorrect operator use (using `@` with a scalar).
        4. STRING SAFETY: Ensure all plot titles/labels using LaTeX (e.g. $\\gamma$) use `r"..."` strings.
        5. SOLVER CRASHES: Ensure 'eps' is NOT passed to CLARABEL.
        6. ATTRIBUTE ERRORS: Ensure `solver.name` is NOT used (replace with `solver`).
        
        If errors exist, return the FULL FIXED code. If it is perfect, return the original code.
        """
        # FIX: Zero temperature for strict checking
        final_resp = model.generate_content(debugger_prompt, generation_config={"temperature": 0.0})
        final_code = extract_code(final_resp.text)
        
        if final_code:
            status.update(label="🕵️ Debugger: Code verified and fixed!", state="complete")
            return final_code
        else:
            status.update(label="🕵️ Debugger: Original code was perfect.", state="complete")
            return raw_code

def run_code_backend(code):
    patch_header = "import matplotlib\nmatplotlib.use('Agg')\nimport matplotlib.pyplot as plt\nimport uuid\ndef mock_show(*args,**kwargs):\n  plt.savefig(f'plot_{uuid.uuid4().hex[:6]}.png')\n  plt.close()\nplt.show = mock_show\n"
    patched_code = patch_header + code
    try:
        proc = subprocess.run([sys.executable, "-c", patched_code], capture_output=True, text=True, timeout=60)
    except: return "", "Execution Timeout", []
    
    plots = []
    for f in glob.glob("*.png"):
        with open(f, "rb") as img: plots.append(img.read())
        os.remove(f)
    return proc.stdout, proc.stderr, plots

# --- 4. MAIN INTERFACE ---

if uploaded_file and ("last_file" not in st.session_state or st.session_state.last_file != uploaded_file.name):
    st.session_state.last_file = uploaded_file.name
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    st.session_state.pdf_images = [Image.open(io.BytesIO(page.get_pixmap(dpi=150).tobytes("png"))) for page in doc]
    st.session_state.show_review = False
    st.rerun()

# --- 3. SESSION STATE INIT (Consolidated) ---
if "extracted_text" not in st.session_state: st.session_state.extracted_text = ""
if "architect_plan" not in st.session_state: st.session_state.architect_plan = ""
if "active_code" not in st.session_state: st.session_state.active_code = ""
if "run_output" not in st.session_state: st.session_state.run_output = None
if "show_review" not in st.session_state: st.session_state.show_review = False

t1, t2 = st.tabs(["📄 Step 1: Extraction & Theory", "🚀 Step 2: Agentic Workbench"])

# --- DEMO PROBLEM (Source of Truth) ---
DEFAULT_PROBLEM_DESCRIPTION = """
**Optimization Problem:**
Minimize scalar nu > 0 subject to LMI constraint.

Dimensions: n=2, m=2, l=2.
Matrices:
A = [[1, 3], [2, 6]]
b = [[5], [10]]
F = [[1, 0], [0, 1]]
E_A = [[0, 1], [0, 0]]
E_b = [[0], [1]]
D = 0 (2x2 zero matrix)

M is a square matrix (l x l):
M = p * I (2x2)

LMI Block Matrix (4x4 blocks):
Row 1: [-I, F*lam, A*x-b, 0]
Row 2: [lam*F.T, -lam*I, 0, lam*D.T*M.T]
Row 3: [(A*x-b).T, 0, -nu, (E_A*x-E_b).T*M.T]
Row 4: [0, M*D*lam, M*(E_A*x-E_b), -lam*I]

Tasks:
1. Plot gamma = sqrt(nu) vs p in [0, 2] step 0.1 for two cases:
   - Curve 1: Optimize over {x, lam, nu}
   - Curve 2: Optimize over {lam, nu} with fixed x=[-1, 2].T
2. Plot x1 vs x2 trajectory for Curve 1.
3. For each p, calculate perturbation Omega* and new LMI feasibility (nu_new).
   - Omega* = (v_x0 @ (E_A@x - E_b + D@v_x0).T) / norm(...)**2
   - v_x0 = inv(inv(lam*I) - F.T@F) @ F.T @ (A@x - b)
   - A* = A + F@Omega*...
   - Check new LMI: [[-I, A*x - b*], [..., -nu_new]] <= 0
4. Print table comparing nu_old vs nu_new and values of Omega*.
"""

if not st.session_state.extracted_text:
    st.session_state.extracted_text = DEFAULT_PROBLEM_DESCRIPTION

with t1:
    if st.session_state.get("pdf_images"):
        if st.button("🔍 Extract Problem Verbatim", type="primary"):
            model = genai.GenerativeModel(selected_model_name)
            resp = model.generate_content(["Extract all matrices and the LMI formula exactly as they appear."] + st.session_state.pdf_images)
            st.session_state.extracted_text = resp.text
            st.session_state.show_review = True
            st.rerun()
    
    problem_desc = st.text_area("Source of Truth:", value=st.session_state.get("extracted_text", ""), height=300)
    st.session_state.extracted_text = problem_desc
    
    SETTINGS_FILE = "settings.json"
    saved_template = ""
    if os.path.exists(SETTINGS_FILE):
        with open(SETTINGS_FILE, "r") as f: saved_template = json.load(f).get("golden_code", "")
    
    with st.expander("⭐ Edit Master Golden Template"):
        template = st.text_area("Code Style Reference:", value=saved_template, height=250)
        if st.button("💾 Save Template"):
            with open(SETTINGS_FILE, "w") as f: json.dump({"golden_code": template}, f)

with t2:
    st.info("Step 1: The Architect analyzes the math. You review it. Step 2: The Coder writes the script.")
    
    if st.button("1️⃣ Run Architect (Generate Plan)", type="primary"):
        if selected_model_name and problem_desc:
            plan = run_architect_agent(problem_desc, selected_model_name)
            st.session_state.architect_plan = plan
            st.rerun()

    # Only show Coder step if a plan exists
    if st.session_state.architect_plan:
        st.subheader("📐 Architect's Plan (Editable)")
        # Display editable text area linked to session state
        edited_plan = st.text_area("Review dimensions and logic before coding:", 
                                   value=st.session_state.architect_plan, 
                                   height=300, 
                                   key="architect_plan_input")
        
        # Sync edits back to session state if changed
        if edited_plan != st.session_state.architect_plan:
            st.session_state.architect_plan = edited_plan

        if st.button("2️⃣ Run Coder (Generate Code)", type="primary"):
            st.session_state.active_code = run_coder_and_debugger(edited_plan, template, selected_model_name)
            st.rerun()

    if st.session_state.get("active_code"):
        st.divider()
        st.subheader("🐍 Generated Code")
        code_input = st.text_area("Final Verified Code:", value=st.session_state.active_code, height=400)
        st.session_state.active_code = code_input

        if st.button("▶️ Run Experiment"):
            with st.spinner("Executing Python..."):
                out, err, plots = run_code_backend(code_input)
                st.session_state.run_output = {"out": out, "err": err, "plots": plots}
        
        if st.session_state.get("run_output"):
            res = st.session_state.run_output
            if res['err']: st.error(res['err'])
            if res['out']: st.info(res['out'])
            for p in res['plots']: st.image(p)

            col1, col2 = st.columns([1, 1])
            with col1:
                if Document and st.button("📄 Save to Word Report"):
                    doc = Document()
                    doc.add_heading("PhD Experiment Results", 0)
                    doc.add_heading("Architect Plan", level=1); doc.add_paragraph(st.session_state.architect_plan)
                    doc.add_heading("Generated Code", level=1); doc.add_paragraph(code_input)
                    doc.add_heading("Solver Logs", level=1); doc.add_paragraph(res['out'])
                    for i, p_bytes in enumerate(res['plots']):
                        p_path = f"tmp_{i}.png"
                        with open(p_path, "wb") as f: f.write(p_bytes)
                        doc.add_picture(p_path, width=Inches(5))
                        os.remove(p_path)
                    doc_io = io.BytesIO(); doc.save(doc_io); doc_io.seek(0)
                    st.download_button("📥 Download Report", doc_io, "report.docx")
            
            with col2:
                if st.button("💾 Save Results to History"):
                    if st.session_state.current_case_name:
                        st.session_state.history[st.session_state.current_case_name] = {
                            "extracted_text": st.session_state.get("extracted_text", ""),
                            "architect_plan": st.session_state.get("architect_plan", ""),
                            "active_code": st.session_state.get("active_code", ""),
                            "run_output": serialize_run_output(st.session_state.get("run_output", None))
                        }
                        with open(HISTORY_FILE, "w") as f: json.dump(st.session_state.history, f)
                        st.toast(f"Saved results to '{st.session_state.current_case_name}'!")
                    else:
                        st.error("Please create or load a case in the sidebar first.")